import os
import json
from docx import Document
from typing import Dict, Optional, Any

class DocumentAnalyzer:
    """Analyse réelle de documents pour Forensics et Malware detection"""
    
    def __init__(self, executor):
        self.executor = executor

    async def analyze_document(self, file_path: str) -> dict:
        """Analyse réelle via exiftool et olevba"""
        if not os.path.exists(file_path):
            return {"error": "Fichier introuvable"}

        # 1. Métadonnées réelles
        meta_out, _ = await self.executor.run_command(f"exiftool -json {file_path}")
        
        # 2. Analyse de macros réelles
        macro_out, _ = await self.executor.run_command(f"olevba --json {file_path}")
        
        try:
            metadata = json.loads(meta_out)[0]
            macros = json.loads(macro_out)
            
            return {
                "metadata": {
                    "author": metadata.get("Author"),
                    "software": metadata.get("Software"),
                    "create_date": metadata.get("CreateDate")
                },
                "macros_detected": len(macros) > 0,
                "suspicious_indicators": [m for m in macros if m.get("type") == "Suspicious"]
            }
        except:
            return {"error": "Erreur lors du parsing des résultats"}

    async def modify_document(
        self,
        source_path: str,
        modifications: Dict[str, str],
        output_path: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Modifie le contenu d'un document Word (.docx).
        
        Args:
            source_path: Chemin vers le document source.
            modifications: Dictionnaire de {texte_ancien: texte_nouveau}.
            output_path: Chemin de sortie (défaut: _modified suffix).
            
        Returns:
            Dictionnaire avec le résultat de la modification.
        """
        if not os.path.exists(source_path):
            return {'error': 'Source file not found'}
        
        file_ext = os.path.splitext(source_path)[1].lower()
        if file_ext != '.docx':
            return {'error': 'Only .docx files are supported for modification.'}

        if not output_path:
            base, ext = os.path.splitext(source_path)
            output_path = f"{base}_modified{ext}"

        try:
            document = Document(source_path)
            modification_count = 0

            # Replace text in paragraphs
            for old_text, new_text in modifications.items():
                for para in document.paragraphs:
                    if old_text in para.text:
                        inline = para.runs
                        for i in range(len(inline)):
                            if old_text in inline[i].text:
                                text = inline[i].text.replace(old_text, new_text)
                                inline[i].text = text
                                modification_count += 1
            
            # Replace text in tables
            for old_text, new_text in modifications.items():
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if old_text in para.text:
                                    inline = para.runs
                                    for i in range(len(inline)):
                                        if old_text in inline[i].text:
                                            text = inline[i].text.replace(old_text, new_text)
                                            inline[i].text = text
                                            modification_count += 1

            document.save(output_path)
            
            return {
                'success': True,
                'source_file': source_path,
                'output_file': output_path,
                'modifications_made': modification_count
            }
        except Exception as e:
            if "No module named 'docx'" in str(e):
                 return {'error': "The 'python-docx' library is required. Please install it using 'pip install python-docx'."}
            return {'error': str(e)}